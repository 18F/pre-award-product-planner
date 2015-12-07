# -*- coding: utf-8 -*-
from flask import Flask, send_from_directory, request, jsonify
from flask_restful import Resource, Api, reqparse
from flask_sqlalchemy import SQLAlchemy

from pprint import pprint

import os
import config

from docx import Document
from models import Agency, RFQ, ContentComponent, Base, Session, engine


# set the project root directory as the static folder, you can set others.
app = Flask(__name__, static_folder='app')
app.config['APP_SETTINGS'] = config.DevelopmentConfig
db = SQLAlchemy(app)
api = Api(app, prefix="/api")

terms_text = "AGILE DEVELOPMENT/AGILE SOFTWARE DEVELOPMENT: A proven commercial methodology for software development that is characterized by incremental and iterative processes where releases are produced in close collaboration with the customer. This process improves investment manageability, lowers risk of project failure, shortens the time to realize value, and allows agencies to better adapt to changing needs.\n\nCONTRACTING OFFICER (CO): The Government official responsible for the execution and administration of contracts on behalf of the Government.\n\nCONTRACTING OFFICER’S REPRESENTATIVE (COR): An individual designated by the Contracting Officer to act as his/her representative to assist in managing the contract. The authorities and limitations of a COR appointment are contained in the written letter of appointment.\n\nDAY: A calendar day unless stated otherwise. If a deliverable is due on a weekend or holiday, the deliverable shall be considered due the next business day.\n\nQUARTER: A quarter will be defined as the first of January through the end of March, first of April through the end of June, first of July through the end of September, and first of October through the end of December.\n\nBUSINESS DAY: Any day other than a Saturday, a Sunday, a Federal holiday or other day on which the Federal Government by law or executive order is closed. Note: This includes any weather-related office closures if the place of performance is in a Federal Building.\n\nMINIMUM FUNCTIONALITY: The minimum capabilities a product should have to meet the Government’s objectives.\n\nAGILE ENVIRONMENT: A team-based setting for IT product development where the Agile development methodology is used.\n\nITERATION/SPRINT/RELEASE CYCLE:Divisions of time within the Agile development framework.  Each iteration is small in scale (i.e., encompasses a single or a few function(s) within a multistep process). Multiple iterations form releases. For more information, see the TechFAR at https://github.com/WhiteHouse/playbook/blob/gh-pages/_includes/techfar-online.md\n\nMILESTONES/EPICS:A necessary step in a process. In this document, used to refer to components of a given project.\n\nSTORY POINT:A measurement of work and effort. Story points are used in an Agile development environment to demonstrate how much work was achieved in a given sprint or iteration. For more information, see the <a href='https://github.com/WhiteHouse/playbook/blob/gh-pages/_includes/techfar-online.md' target='_blank'>TechFAR</a>\n\nTHROUGHPUT: The amount of material or items passing through a system or process; in this document, refers to the work activity of a product development team.";
payment_schedule_text = "The contractor shall be paid upon the completion of each iteration upon its acceptance and verification by the Contracting Officer’s Representative (COR). Invoices shall be submitted at the end of each iteration in accordance with the delivery schedule as established in the Performance Work Statement."

DATA = {
    "definitions": terms_text, 
    "payment_schedule": payment_schedule_text,
}

# def abort_if_content_doesnt_exist(content_key):
#     if content_key not in DATA:
#         abort(404, message="Content {} doesn't exist".format(key))

parser = reqparse.RequestParser()
parser.add_argument('agency')
parser.add_argument('doc_type')
parser.add_argument('setaside')

class Agencies(Resource):
    def get(self):
        agencies = session.query(Agency).order_by(Agency.full_name).all()
        return jsonify(agencies)

class Data(Resource):

    def get(self, content_key):        
        print "content key! " + content_key
        session = Session()
        content = session.query(ContentComponent).filter_by(name=content_key)
        return jsonify(content)

    def put(self, content_key):
        data = request.get_json()
        content = data['text']
        print "content '" + content + "'"
        DATA[content_key] = content

class Create(Resource):

     def post(self):
        # get agency, doc_type, setaside values
        args = parser.parse_args()
        agency = args['agency']
        doc_type = args['doc_type']
        setaside = args['setaside']

        rfq = RFQ(agency=agency, doc_type=doc_type, setaside=setaside)
        session.add(rfq)
        session.commit()


api.add_resource(Data, '/get_content/<string:content_key>')
api.add_resource(Create, '/rfqs')
api.add_resource(Agencies, '/agencies')

# map index.html to app/index.html, map /build/bundle.js to app/build.bundle.js
@app.route('/initiate')
def initiate():
    RFQ.create(agency="", agency_full_name="", doc_type="")

@app.route('/')
def index():
    return send_from_directory("app", "index.html")

@app.route('/<path:path>')
def send_js(path):
    return send_from_directory("app", path)

@app.route('/download/<int:doc_id>')
def download():
    document = Document()

    rfq = RFQ.find(id=doc_id)
    title = "RFQ for " + rfq.agency
    document.add_heading('RFQ', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='IntenseQuote')

    document.add_paragraph(
        'first item in unordered list', style='ListBullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='ListNumber'
    )
    document.save('demo.docx')

    return "document created"

if __name__ == "__main__":
    app.run(debug=True)
